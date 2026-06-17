import os
import gc
import tempfile
import torch
import gradio as gr
import whisperx
from whisperx.diarize import DiarizationPipeline
from dotenv import load_dotenv, set_key
import argostranslate.package
import argostranslate.translate
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt

env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
load_dotenv(env_path)

def format_timestamp(seconds: float) -> str:
    """Convierte segundos a formato de tiempo SRT: HH:MM:SS,mmm"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    milliseconds = int(round((seconds - int(seconds)) * 1000))
    # Asegurar límites de milisegundos
    milliseconds = min(999, max(0, milliseconds))
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

def transcribe_and_diarize(file_path, hf_token, whisper_model, source_language_ui, min_speakers, max_speakers, traducir_a_es, progress=gr.Progress(track_tqdm=True)):
    # Validaciones iniciales
    if not file_path:
        return "Error: Por favor, carga un archivo de audio o video.", None, None, gr.update(), None
    if not hf_token or not hf_token.strip():
        return "Error: El token de Hugging Face es obligatorio para descargar los modelos de Diarización (Pyannote).", None, None, gr.update(), None
        
    # Auto-guardado del token en .env para la próxima vez
    current_token = os.getenv("HF_TOKEN", "")
    if hf_token.strip() != current_token:
        try:
            if not os.path.exists(env_path):
                open(env_path, 'w').close()
            set_key(env_path, "HF_TOKEN", hf_token.strip())
            os.environ["HF_TOKEN"] = hf_token.strip()
            print("[INFO] Token de Hugging Face guardado localmente en .env.")
        except Exception as e:
            print(f"[WARNING] No se pudo guardar el token en .env: {e}")
            
    device = "cuda" if torch.cuda.is_available() else "cpu"
    if device == "cpu":
        return "Error: No se detectó una GPU compatible con CUDA. Esta aplicación requiere aceleración por GPU (NVIDIA) para funcionar.", None, None, gr.update(), None
        
    compute_type = "int8"  # int8 evita errores de kernel ('no kernel image') en GPUs más antiguas y optimiza la VRAM.
    batch_size = 4
    
    temp_txt_path = None
    temp_srt_path = None
    
    try:
        # --- PASO 1: Transcripción con WhisperX ---
        progress(0.1, desc="Cargando modelo WhisperX en VRAM...")
        print(f"[PROCESS] Cargando modelo WhisperX: {whisper_model}...")
        model = whisperx.load_model(whisper_model, device, compute_type=compute_type)
        
        progress(0.2, desc="Procesando audio de entrada...")
        print("[PROCESS] Cargando audio...")
        audio = whisperx.load_audio(file_path)
        
        progress(0.3, desc="Ejecutando transcripción de audio...")
        print("[PROCESS] Iniciando transcripción...")
        
        # Mapeo de idioma de la UI a código Whisper
        lang_map = {
            "Español": "es", "Inglés": "en", "Francés": "fr", 
            "Alemán": "de", "Italiano": "it", "Portugués": "pt"
        }
        language_code = lang_map.get(source_language_ui, None)
        
        if language_code:
            raw_result = model.transcribe(audio, batch_size=batch_size, language=language_code)
        else:
            raw_result = model.transcribe(audio, batch_size=batch_size)
        
        # Guardar idioma detectado para cargar el modelo de alineación correspondiente
        detected_language = raw_result.get("language", "es")
        print(f"[PROCESS] Idioma detectado: {detected_language}")
        
        # Limpieza de memoria (VRAM) inmediata
        progress(0.4, desc="Liberando VRAM de transcripción...")
        print("[PROCESS] Liberando memoria de transcripción...")
        del model
        gc.collect()
        torch.cuda.empty_cache()
        
        # --- PASO 2: Alineación fina (Wav2Vec2) ---
        progress(0.5, desc="Cargando modelo de alineación fonética...")
        print("[PROCESS] Cargando modelo de alineación...")
        align_model, metadata = whisperx.load_align_model(language_code=detected_language, device=device)
        
        progress(0.6, desc="Alineando palabras con precisión de milisegundos...")
        print("[PROCESS] Ejecutando alineación...")
        aligned_result = whisperx.align(raw_result["segments"], align_model, metadata, audio, device, return_char_alignments=False)
        
        # Limpieza de memoria (VRAM) inmediata
        progress(0.7, desc="Liberando VRAM de alineación...")
        print("[PROCESS] Liberando memoria de alineación...")
        del align_model
        gc.collect()
        torch.cuda.empty_cache()
        
        # --- PASO 3: Diarización (Pyannote) ---
        progress(0.8, desc="Iniciando Diarización de hablantes (Pyannote)...")
        print("[PROCESS] Inicializando pipeline de Diarización...")
        
        # Configurar los parámetros opcionales de cantidad de hablantes
        diarize_kwargs = {}
        if min_speakers is not None and min_speakers > 0:
            diarize_kwargs["min_speakers"] = int(min_speakers)
        if max_speakers is not None and max_speakers > 0:
            diarize_kwargs["max_speakers"] = int(max_speakers)
            
        diarize_model = DiarizationPipeline(token=hf_token.strip(), device=device)
        
        progress(0.85, desc="Separando pistas de voces y hablantes...")
        print("[PROCESS] Ejecutando diarización...")
        diarize_segments = diarize_model(audio, **diarize_kwargs)
        
        # Limpieza de memoria (VRAM) inmediata
        progress(0.9, desc="Liberando VRAM de diarización...")
        print("[PROCESS] Liberando memoria de diarización...")
        del diarize_model
        gc.collect()
        torch.cuda.empty_cache()
        
        # --- PASO 4: Fusión final de palabras y hablantes ---
        progress(0.95, desc="Fusionando texto con etiquetas de hablantes...")
        print("[PROCESS] Asignando hablantes a palabras...")
        final_result = whisperx.assign_word_speakers(diarize_segments, aligned_result)
        
        # --- PASO 4.5: Traducción (Opcional) ---
        if traducir_a_es and detected_language != "es":
            progress(0.96, desc="Inicializando motor de Traducción (Argos)...")
            print("[PROCESS] Preparando traducción a Español...")
            try:
                argostranslate.package.update_package_index()
                available_packages = argostranslate.package.get_available_packages()
                
                # Buscar el paquete adecuado (Origen -> Español)
                package_to_install = next(
                    filter(
                        lambda x: x.from_code == detected_language and x.to_code == "es", available_packages
                    ), None
                )
                
                if package_to_install is not None:
                    progress(0.97, desc=f"Descargando paquete de idioma ({detected_language} -> es) si es necesario...")
                    package_to_install.download()
                    package_to_install.install()
                    
                    installed_languages = argostranslate.translate.get_installed_languages()
                    from_lang = list(filter(lambda x: x.code == detected_language, installed_languages))[0]
                    to_lang = list(filter(lambda x: x.code == "es", installed_languages))[0]
                    translation_obj = from_lang.get_translation(to_lang)
                    
                    progress(0.98, desc="Traduciendo transcripción al Español...")
                    print("[PROCESS] Traducción en progreso...")
                    for segment in final_result["segments"]:
                        original_text = segment.get("text", "").strip()
                        if original_text:
                            segment["text"] = translation_obj.translate(original_text)
                else:
                    print(f"[WARNING] No se encontró paquete de traducción directo de '{detected_language}' a 'es'.")
            except Exception as tr_e:
                print(f"[ERROR] Falló la traducción: {tr_e}")

        # --- PASO 5: Generar y formatear salidas ---
        progress(0.99, desc="Formateando subtítulos finales...")
        df_rows = []
        srt_lines = []
        srt_counter = 1
        
        for segment in final_result["segments"]:
            speaker = segment.get("speaker", "SPEAKER_UNKNOWN")
            text = segment.get("text", "").strip()
            start = segment.get("start", 0.0)
            end = segment.get("end", 0.0)
            
            # Fila para el Dataframe
            df_rows.append({"Hablante": speaker, "Inicio (s)": round(start, 2), "Fin (s)": round(end, 2), "Texto": text})
            
            # Formato de archivo SRT
            srt_lines.append(f"{srt_counter}")
            srt_lines.append(f"{format_timestamp(start)} --> {format_timestamp(end)}")
            srt_lines.append(f"[{speaker}]: {text}\n")
            srt_counter += 1
            
        df_data = pd.DataFrame(df_rows)
        full_srt = "\n".join(srt_lines)
        
        if df_data.empty:
            df_data = pd.DataFrame([{"Hablante": "Sistema", "Inicio (s)": 0.0, "Fin (s)": 0.0, "Texto": "No se detectó contenido de voz en el archivo."}])
            
        # Generar los archivos temporales de descarga iniciales
        temp_txt_path, temp_srt_path, temp_docx_path = update_downloads_from_df(df_data, clean_export=False)
            
        progress(1.0, desc="¡Proceso finalizado!")
        print("[SUCCESS] Proceso de diarización y transcripción completado.")
        return df_data, temp_txt_path, temp_srt_path, temp_docx_path, gr.Tabs(selected=2), (file_path, temp_srt_path)
        
    except Exception as e:
        error_msg = f"ERROR EN EL PROCESAMIENTO:\n\n{str(e)}\n\n"
        error_msg += "Sugerencia: Asegúrate de que el token de Hugging Face sea válido y de haber aceptado las licencias de Pyannote en su web."
        print(f"[ERROR] {error_msg}")
        
        # Liberar memoria VRAM en caso de error crítico
        try:
            gc.collect()
            torch.cuda.empty_cache()
        except:
            pass
            
        error_df = pd.DataFrame([{"Hablante": "ERROR", "Inicio (s)": 0.0, "Fin (s)": 0.0, "Texto": error_msg}])
        return error_df, None, None, None, gr.update(), None

def get_speaker_color(speaker_name):
    hash_val = sum(ord(c) for c in speaker_name)
    colors = [
        RGBColor(0, 114, 178), # Blue
        RGBColor(213, 94, 0),  # Vermillion
        RGBColor(0, 158, 115), # Bluish Green
        RGBColor(204, 121, 167), # Reddish Purple
        RGBColor(230, 159, 0), # Orange
        RGBColor(86, 180, 233), # Sky Blue
    ]
    return colors[hash_val % len(colors)]

def update_downloads_from_df(df, clean_export=False):
    if df is None or df.empty:
        return None, None, None
    
    srt_lines = []
    txt_lines = []
    srt_counter = 1
    
    doc = Document()
    doc.add_heading('Guion de Transcripción', 0)
    
    for _, row in df.iterrows():
        speaker = str(row.get("Hablante", "SPEAKER_UNKNOWN"))
        try:
            start = float(row.get("Inicio (s)", 0.0))
            end = float(row.get("Fin (s)", 0.0))
        except:
            start, end = 0.0, 0.0
        text = str(row.get("Texto", "")).strip()
        
        # DOCX Logic
        p = doc.add_paragraph()
        speaker_run = p.add_run(f"[{speaker}]")
        speaker_run.bold = True
        speaker_run.font.color.rgb = get_speaker_color(speaker)
        p.add_run(f"\n({start:.2f}s - {end:.2f}s) {text}")
        
        if clean_export:
            txt_lines.append(f"({start:.2f}s - {end:.2f}s): {text}")
            
            srt_lines.append(f"{srt_counter}")
            srt_lines.append(f"{format_timestamp(start)} --> {format_timestamp(end)}")
            srt_lines.append(f"{text}\n")
        else:
            txt_lines.append(f"[{speaker}] ({start:.2f}s - {end:.2f}s): {text}")
            
            srt_lines.append(f"{srt_counter}")
            srt_lines.append(f"{format_timestamp(start)} --> {format_timestamp(end)}")
            srt_lines.append(f"[{speaker}]: {text}\n")
            
        srt_counter += 1
        
    full_text = "\n".join(txt_lines)
    full_srt = "\n".join(srt_lines)
    
    temp_dir = tempfile.gettempdir()
    temp_txt_path = os.path.join(temp_dir, "transcripcion_diarizada_editada.txt")
    temp_srt_path = os.path.join(temp_dir, "transcripcion_diarizada_editada.srt")
    temp_docx_path = os.path.join(temp_dir, "transcripcion_guion.docx")
    
    with open(temp_txt_path, "w", encoding="utf-8") as f:
        f.write(full_text)
        
    with open(temp_srt_path, "w", encoding="utf-8") as f:
        f.write(full_srt)
        
    doc.save(temp_docx_path)
        
    return temp_txt_path, temp_srt_path, temp_docx_path

# --- CSS Personalizado para una Estética Premium Oscura ---
custom_css = """
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');

body, .gradio-container {
    font-family: 'Outfit', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    background: radial-gradient(circle at top, #1e1b4b 0%, #030712 100%) !important;
}

/* Panel principal con Glassmorphism */
.glass-panel {
    background: rgba(17, 24, 39, 0.7) !important;
    backdrop-filter: blur(16px) !important;
    border: 1px solid rgba(255, 255, 255, 0.08) !important;
    border-radius: 20px !important;
    box-shadow: 0 10px 40px 0 rgba(0, 0, 0, 0.5) !important;
    padding: 25px !important;
}

/* Header estilizado */
.app-header {
    text-align: center;
    padding: 2rem 1rem;
    margin-bottom: 2rem;
    background: linear-gradient(135deg, rgba(99, 102, 241, 0.1) 0%, rgba(168, 85, 247, 0.1) 100%);
    border: 1px solid rgba(99, 102, 241, 0.15);
    border-radius: 16px;
    box-shadow: 0 4px 30px rgba(0, 0, 0, 0.3);
}

.app-title {
    font-size: 2.6rem;
    font-weight: 700;
    margin: 0;
    background: linear-gradient(90deg, #818cf8, #c084fc, #f472b6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    letter-spacing: -0.03em;
}

.app-subtitle {
    font-size: 1.05rem;
    color: #94a3b8;
    margin-top: 0.5rem;
    font-weight: 400;
}

/* Botones de acción principal */
.action-btn {
    background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%) !important;
    border: none !important;
    color: white !important;
    font-weight: 600 !important;
    border-radius: 12px !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3) !important;
}

.action-btn:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 25px rgba(99, 102, 241, 0.5) !important;
}

/* Modificaciones de cajas de Gradio */
input, textarea, select {
    transition: border-color 0.2s, box-shadow 0.2s !important;
}

input:focus, textarea:focus, select:focus {
    border-color: #818cf8 !important;
    box-shadow: 0 0 0 3px rgba(129, 140, 248, 0.25) !important;
}
"""

# --- CONSTRUCCIÓN DE LA INTERFAZ DE USUARIO ---
with gr.Blocks(theme=gr.themes.Default(primary_hue="indigo", secondary_hue="slate", neutral_hue="zinc"), css=custom_css) as demo:
    
    # Encabezado principal
    gr.HTML(
        """
        <div class="app-header">
            <h1 class="app-title">EchoWriter AI</h1>
            <p class="app-subtitle">Transcripción Local Inteligente & Separación de Hablantes (Diarización)</p>
        </div>
        """
    )
    
    with gr.Tabs() as main_tabs:
        with gr.TabItem("⚙️ 1. Configuración de Entrada", id=1):
            with gr.Column(elem_classes="glass-panel"):
                gr.Markdown("### 🛠️ Carga de Archivo Multimedia")
                
                audio_video_input = gr.File(
                    label="Subir o arrastrar Archivo (Audio/Video)",
                    file_types=["audio", "video"],
                    type="filepath"
                )
                
                with gr.Accordion("⚙️ Parámetros del Motor de IA", open=True):
                    hf_token_input = gr.Textbox(
                        label="Token de Hugging Face (pyannote)",
                        placeholder="hf_...",
                        type="password",
                        value=os.getenv("HF_TOKEN", ""),
                        info="Requerido para la diarización. Asegúrate de haber aceptado la licencia de Pyannote 3.1 en Hugging Face."
                    )
                    
                    with gr.Row():
                        whisper_model_input = gr.Dropdown(
                            label="Modelo de Transcripción Whisper",
                            choices=["small", "medium", "turbo"],
                            value="turbo",
                            info="El modelo 'large' está deshabilitado para evitar errores (OOM)."
                        )
                        
                        source_language_input = gr.Dropdown(
                            label="Idioma del Audio (Opcional)",
                            choices=["Automático", "Español", "Inglés", "Francés", "Alemán", "Italiano", "Portugués"],
                            value="Automático",
                            info="Seleccionar el idioma manualmente salta el paso de auto-detección."
                        )
                    
                    with gr.Row():
                        min_speakers_input = gr.Number(
                            label="Min Hablantes (Opcional)", 
                            precision=0, 
                            minimum=0,
                            value=0
                        )
                        max_speakers_input = gr.Number(
                            label="Max Hablantes (Opcional)", 
                            precision=0, 
                            minimum=0,
                            value=0
                        )
                        
                    traducir_es_input = gr.Checkbox(
                        label="Traducir transcripción al Español (Offline)",
                        value=False,
                        info="Si el audio está en otro idioma, la IA lo traducirá automáticamente al español."
                    )
                    
                submit_btn = gr.Button("Comenzar Transcripción", elem_classes="action-btn")
        
        with gr.TabItem("🎬 2. Edición y Auditoría", id=2):
            with gr.Column(elem_classes="glass-panel"):
                gr.Markdown("### 📝 Auditoría de Transcripción (Sincronización)")
                
                with gr.Row():
                    # Columna Izquierda: Reproductor
                    with gr.Column(scale=1):
                        playback_video = gr.Video(
                            label="Reproductor de Auditoría",
                            interactive=False
                        )
                        
                    # Columna Derecha: Texto
                    with gr.Column(scale=1):
                        output_dataframe = gr.Dataframe(
                            headers=["Hablante", "Inicio (s)", "Fin (s)", "Texto"],
                            datatype=["str", "number", "number", "str"],
                            label="Transcripción Editable (Doble clic en una celda para editar)",
                            interactive=True,
                            wrap=True
                        )
                        
                        clean_export_input = gr.Checkbox(
                            label="🧹 Modo Subtítulos Limpios (Ocultar nombres de hablantes en descarga)",
                            value=False,
                            info="Útil para exportar a Premiere, YouTube o DaVinci Resolve sin etiquetas molestas."
                        )
                        
                        apply_edits_btn = gr.Button("💾 Aplicar Ediciones y Preparar Descargas", elem_classes="action-btn")
                        
                        with gr.Row():
                            txt_download_btn = gr.File(
                                label="Descargar Transcripción (.txt)",
                                interactive=False
                            )
                            srt_download_btn = gr.File(
                                label="Descargar Subtítulos (.srt)",
                                interactive=False
                            )
                            docx_download_btn = gr.File(
                                label="📄 Descargar Guion Profesional (.docx)",
                                interactive=False
                            )

    # Vinculación de eventos
    submit_btn.click(
        fn=transcribe_and_diarize,
        inputs=[
            audio_video_input,
            hf_token_input,
            whisper_model_input,
            source_language_input,
            min_speakers_input,
            max_speakers_input,
            traducir_es_input
        ],
        outputs=[
            output_dataframe,
            txt_download_btn,
            srt_download_btn,
            docx_download_btn,
            main_tabs,
            playback_video
        ]
    )
    
    apply_edits_btn.click(
        fn=update_downloads_from_df,
        inputs=[output_dataframe, clean_export_input],
        outputs=[txt_download_btn, srt_download_btn, docx_download_btn]
    )

# Lanzar aplicación localmente
if __name__ == "__main__":
    # La aplicación corre localmente en el puerto 7860
    demo.launch(server_name="127.0.0.1", server_port=7860, share=False)
